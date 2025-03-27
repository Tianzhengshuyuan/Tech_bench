import argparse
from docx import Document

def replace_line_breaks_with_paragraphs(file_path, output_path):
    # 打开 Word 文档
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        # 替换每个段落中的换行符（\n）为段落结束符
        if '\n' in paragraph.text:
            lines = paragraph.text.split('\n')
            for i, line in enumerate(lines):
                if i == 0:
                    paragraph.text = line
                else:
                    # 添加新段落
                    doc.add_paragraph(line)

    # 保存修改后的文档
    doc.save(output_path)

# 命令行入口
if __name__ == "__main__":
    # 定义命令行参数
    parser = argparse.ArgumentParser(description="去除word文档中存在的仅换行不开始新段落的\n，方便解析")
    parser.add_argument("--docx_name", type=str, help="Word文档名")
    args = parser.parse_args()
    
    docx_path = f"./{args.docx_name}.docx"  
    
    # 处理文档
    replace_line_breaks_with_paragraphs(docx_path, docx_path)
    
    print(f"完成\n移除，文件已保存为 {docx_path}")