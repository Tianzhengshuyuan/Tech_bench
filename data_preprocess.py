import re
import json
import pytesseract
import zipfile
import subprocess
import requests
import os
import argparse
from docx import Document
from wand.image import Image
from PIL import Image as PILImage
from io import BytesIO
from xml.etree.ElementTree import Element
from xml.etree import ElementTree as ET

def parse_relationships(docx_zip):
    """
    解析 document.xml.rels 文件，建立 rId 与 Target 文件路径的映射关系。
    """
    relationships = {}
    with docx_zip.open('word/_rels/document.xml.rels') as rels_file:
        # 解析 XML
        tree = ET.parse(rels_file)
        root = tree.getroot()
        for rel in root:
            rId = rel.attrib.get("Id")
            target = rel.attrib.get("Target")
            if rId and target:
                relationships[rId] = target
    return relationships

def extract_high_res_image_from_docx(docx_path, rId, relationships):
    """
    从解压后的 Word 文件中提取高分辨率 WMF 图像，并转换为 PNG。
    """
    # 获取 rId 对应的 Target 路径
    target_path = relationships.get(rId)
    print(target_path)
    if not target_path:
        print(f"No target found for rId: {rId}")
        return None

    # 确保路径是相对路径，拼接完整路径
    if not target_path.startswith('/'):
        target_path = f"word/{target_path}"

    # 从 ZIP 文件中提取目标文件
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        if target_path in docx_zip.namelist():
            with docx_zip.open(target_path) as image_file:
                image_data = image_file.read()
                wmf_path = f"./wmf_images/{rId}.wmf"
                with open(wmf_path, "wb") as wmf_file:
                    wmf_file.write(image_data)
                output_path = f"./png_images/{rId}.png"
                output_dir = f"./png_images/"
                # 构造命令
                command = [
                    "libreoffice", "--headless", "--convert-to", "png", wmf_path, "--outdir", output_dir
                ]
                
                # 执行命令
                result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                
                return PILImage.open(output_path)
        else:
            print(f"Target path {target_path} not found in ZIP.")
    return None

def extract_text_or_formula(run, dotx_path, relationships):
    """
    提取段落中的文本或公式图片并识别为 LaTeX。
    """
    run_xml = run.element  # 获取当前运行对象的 XML 元素
    print(run_xml.xml)
    if run.text.strip():
        # 如果是纯文本，直接返回
        return run.text.strip()

    elif "<w:object" in run_xml.xml:
        # 解析 OLE 对象
        for ole_object in run_xml.findall(".//{urn:schemas-microsoft-com:office:office}OLEObject"):
            prog_id = ole_object.attrib.get("ProgID", "")

            # 检查是否是公式
            if prog_id in ["Equation.3", "MathType"]:
                rId = ole_object.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                # 使用正则表达式匹配 "rId" 后的数字部分
                match = re.match(r"(rId)(\d+)", rId)
                if match:
                    prefix = match.group(1)  # 匹配 "rId"
                    number = int(match.group(2))  # 提取数字部分并转换为整数
                    number -= 1  # 修改数字
                    rId = f"{prefix}{number}"  # 拼接新的字符串
                    
                if rId:
                    # 使用 rId 提取高分辨率图像
                    img = extract_high_res_image_from_docx(docx_path, rId, relationships)
                    file_path = f"./png_images/{rId}.png"
                    # if img:
                    #     # 使用 SimpleTex 的 API 识别公式
                    #     SIMPLETEX_UAT="x97YHMaxT4hl1kbcvKkAHbQqZGR0HDL0rBAZfmqScLusUcO74sXCCOIsNfqO3PgM"
                    #     api_url="https://server.simpletex.cn/api/latex_ocr"  # 接口地址
                    #     data = { } # 请求参数数据（非文件型参数），视情况填入，可以参考各个接口的参数说明
                    #     header={ "token": SIMPLETEX_UAT } # 鉴权信息，此处使用UAT方式
                    #     file=[("file",(file_path,open(file_path, 'rb')))] # 请求文件,字段名一般为file
                    #     res = requests.post(api_url, files=file, data=data, headers=header) # 使用requests库上传文件
                    #     content = json.loads(res.text)['res']['latex']
                    #     return content
                    # else:
                    #     print(f"Could not find image for rId: {rId}")
    return ""

def extract_questions_from_docx(docx_path, output_json_path):
    # 打开 Word 文档
    doc = Document(docx_path)
    
    # 解压文档，解析 rId 和 Target 的映射关系
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        relationships = parse_relationships(docx_zip)

    # 定义正则表达式，用于匹配选择题的格式
    # (?P<question>...)：定义一个命名捕获组，名为 question，用于捕获匹配的内容。
    # \d+：匹配一个或多个数字，表示题目的编号
    # \.：匹配一个点号（.），表示题号和题目内容之间的分隔符。
    # \n：匹配换行符，表示题目内容结束。
    # .*?：非贪婪匹配任意字符（除了换行符），表示题目内容。
    # \s*：匹配零个或多个空白字符（如空格、制表符等），允许在 A. 和选项内容之间有任意数量的空白字符。

    question_pattern = re.compile(
        r"(?P<question>\d+[．.。].*?)"           # 匹配题目开头，如 "17. 一定质量的..."
        r"(?:A[．.。]\s*(?P<A>.*?))"             # 匹配 A 选项
        r"(?:B[．.。]\s*(?P<B>.*?))"             # 匹配 B 选项
        r"(?:C[．.。]\s*(?P<C>.*?))"             # 匹配 C 选项
        r"(?:D[．.。]\s*(?P<D>.*?))"             # 匹配 D 选项
        r"(?=\n|\f|\d+[．.。][\u4e00-\u9fa5])",     # 断言 D 选项后面是换行符、换页符、题号或中文汉字，但不包含这些内容
        re.DOTALL                           # 允许匹配跨行内容
    )

    # 提取文档中的文本
    full_text = "\n".join([p.text for p in doc.paragraphs])


    # 提取选择题内容
    questions = []
    for match in question_pattern.finditer(full_text):
        print(match.group(0))
        question_data = match.groupdict()
        question_data['index'] = (re.match(r'^(\d+)[．.]', question_data["question"][:10])).group(1)
        for i,paragraph in enumerate(doc.paragraphs):
            if paragraph.text.startswith(question_data["question"][:10].strip()):
                #找到option_paragraph的真正起点
                start = 1
                for j in range(1,10):
                    if doc.paragraphs[i+j].text.startswith("A"):
                        start = i+j
                        break
                option_paragraph1 = doc.paragraphs[start]
                option_paragraph2 = doc.paragraphs[start+1]
                option_paragraph3 = doc.paragraphs[start+2]
                option_paragraph4 = doc.paragraphs[start+3]
                results_A = ""
                results_B = ""
                results_C = ""
                results_D = ""
                print(option_paragraph1.text)
                #四个选项各占一行
                if option_paragraph1.text.startswith("A") and option_paragraph2.text.startswith("B"): 
                    print("situation1")
                    option_count=0
                    for i,run in enumerate(option_paragraph1.runs):
                        result = extract_text_or_formula(run, docx_path, relationships).lstrip("．.。")
                        if re.search(r"A", result): 
                            option_count+=1
                            if len(result) > 2: #考虑到有时 “C．甲、丁的种群数量下降，丙的种群数量增加”会被解析为一个完整的run
                                results_A += result[2:]
                        elif option_count==1:
                            results_A += result
                    for i,run in enumerate(option_paragraph2.runs):
                        result = extract_text_or_formula(run, docx_path, relationships).lstrip("．.。")
                        if re.search(r"B", result):
                            option_count+=1
                            if len(result) > 2:
                                results_B += result[2:]
                        elif option_count==2:
                            results_B += result
                    for i,run in enumerate(option_paragraph3.runs):
                        result = extract_text_or_formula(run, docx_path, relationships).lstrip("．.。")
                        if re.search(r"C", result):
                            option_count+=1
                            if len(result) > 2:
                                results_C += result[2:]
                        elif option_count==3:
                            results_C += result
                    for i,run in enumerate(option_paragraph4.runs):
                        result = extract_text_or_formula(run, docx_path, relationships).lstrip("．.。")
                        if re.search(r"D", result):
                            option_count+=1
                            if len(result) > 2:
                                results_D += result[2:]
                        elif option_count==4:
                            results_D += result
                #四个选项占两行
                elif option_paragraph1.text.startswith("A") and option_paragraph2.text.startswith("C"):
                    print("situation2") 
                    option_count=0
                    for i,run in enumerate(option_paragraph1.runs):
                        result = extract_text_or_formula(run, docx_path, relationships).lstrip("．.。")
                        if result.startswith("A") or result.startswith("B"):
                            option_count+=1
                            if len(result) > 2:
                                if option_count==1:
                                    results_A += result[2:]
                                else:
                                    results_B += result[2:]
                        else:
                            if result.endswith(("A", "B")): #可能遇到 <w:t>＝25cm/s，向左传播         B．</w:t>
                                if option_count==1:
                                    results_A += result[:-1]
                                elif option_count==2:
                                    results_B += result[:-1]  
                                option_count += 1
                            elif result.endswith(("A.", "B.", "A．", "B．")):
                                if option_count==1:
                                    results_A += result[:-2]
                                elif option_count==2:
                                    results_B += result[:-2]                               
                                option_count += 1
                            else:
                                if option_count==1:
                                    results_A += result
                                elif option_count==2:
                                    results_B += result
                    for i,run in enumerate(option_paragraph2.runs):
                        result = extract_text_or_formula(run, docx_path, relationships).lstrip("．.。")
                        if result.startswith("C") or result.startswith("D"):
                            option_count+=1
                            if len(result) > 2:
                                if option_count==3:
                                    results_C += result[2:]
                                else:
                                    results_D += result[2:]
                        else:
                            if result.endswith(("C", "D")): 
                                if option_count==3:
                                    results_C += result[:-1]
                                elif option_count==4:
                                    results_D += result[:-1]  
                                option_count += 1
                            elif result.endswith(("C.", "D.", "C．", "D．")):
                                if option_count==3:
                                    results_C += result[:-2]
                                elif option_count==4:
                                    results_D += result[:-2]                               
                                option_count += 1
                            else:
                                if option_count==3:
                                    results_C += result
                                elif option_count==4:
                                    results_D += result                        
                else: #四个选项在同一行
                    print("situation3")
                    option_count=0
                    for i,run in enumerate(option_paragraph1.runs):
                        result = extract_text_or_formula(run, docx_path, relationships).lstrip("．.。")
                        if re.search(r"(^|[^a-zA-Z])A([^a-zA-Z]|$)", result) or re.search(r"(^|[^a-zA-Z])B([^a-zA-Z]|$)", result) or re.search(r"(^|[^a-zA-Z])C([^a-zA-Z]|$)", result) or re.search(r"(^|[^a-zA-Z])D([^a-zA-Z]|$)", result):
                            #全是文本，一整行被解析成一个run
                            if re.search(r"A", result) and re.search(r"B", result) and re.search(r"C", result) and re.search(r"D", result):
                                # 定义正则表达式匹配整个模式
                                pattern = r"A[．.](.*?)B[．.](.*?)C[．.](.*?)D[．.](.*)"

                                # 使用 re.search 寻找第一次匹配
                                match_option = re.search(pattern, result)

                                # 提取各部分内容并去掉多余空白
                                if match_option:
                                    results_A = match_option.group(1).strip()  # A. 和 B. 之间的内容
                                    results_B = match_option.group(2).strip()  # B. 和 C. 之间的内容
                                    results_C = match_option.group(3).strip()  # C. 和 D. 之间的内容
                                    results_D = match_option.group(4).strip()  # D. 后面的内容
                            else:
                                option_count += 1
                                # print(result)
                                # print("result len is: "+str(len(result)))
                                if len(result) > 2 and result.startswith(("A", "B", "C", "D")):
                                    print("hello")
                                    if result.endswith(("B", "C", "D")): #考虑到可能出现 <w:t xml:space="preserve">    B．0        C．</w:t>
                                        if option_count==1:
                                            results_A += result[2:-1]
                                        elif option_count==2:
                                            results_B += result[2:-1]
                                        elif option_count==3:
                                            results_C += result[2:-1]
                                        else:
                                            results_D += result[2:-1]  
                                        option_count += 1 
                                    elif result.endswith(("B.", "C.", "D.", "B．", "C．", "D．")):  
                                        if option_count==1:
                                            results_A += result[2:-2]
                                        elif option_count==2:
                                            results_B += result[2:-2]
                                        elif option_count==3:
                                            results_C += result[2:-2]
                                        else:
                                            results_D += result[2:-2]  
                                        option_count += 1                                    
                                    else:
                                        if option_count==1:
                                            results_A += result[2:]
                                        elif option_count==2:
                                            results_B += result[2:]
                                        elif option_count==3:
                                            results_C += result[2:]
                                        else:
                                            results_D += result[2:]
                        elif option_count==1:
                            results_A += result
                        elif option_count==2:
                            results_B += result
                        elif option_count==3:
                            results_C += result
                        else:
                            results_D += result
                question_data["A"] = results_A
                question_data["B"] = results_B
                question_data["C"] = results_C
                question_data["D"] = results_D

        print(">>>>>>>>>>")  
            
        # 添加识别到的问题
        questions.append(question_data)

    # 保存为 JSON 文件
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(questions, f, ensure_ascii=False, indent=4)

    print(f"提取完成！选择题已保存到 {output_json_path}")

# 命令行入口
if __name__ == "__main__":
    # 定义命令行参数
    parser = argparse.ArgumentParser(description="从 Word 文档中提取选择题和答案并保存为 JSON 格式")
    parser.add_argument("--docx_name", type=str, help="Word文档名")
    parser.add_argument("--output", type=str, default="questions.json", help="输出 JSON 文件路径 (默认: questions.json)")
    args = parser.parse_args()
    docx_path = f"/root/tech_bench/docx/{args.docx_name}.docx"  
    output_json_path = f"/root/tech_bench/json/{args.docx_name}.json"     
    extract_questions_from_docx(docx_path, output_json_path)