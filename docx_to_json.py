import re
import os
import json
import zipfile
import requests
import argparse
import subprocess
import pytesseract
import numpy as np
from lxml import etree
from io import BytesIO
from docx import Document
from wand.image import Image
from xml.etree import ElementTree as ET
from xml.etree.ElementTree import Element
from PIL import Image as PILImage, ImageOps

def crop_image(image):
    """
    裁剪图像，去除空白区域，只保留公式部分。
    """
    # 转为灰度图像
    gray_image = ImageOps.grayscale(image)

    # 将灰度图像转换为 NumPy 数组
    image_array = np.array(gray_image)

    # 设置阈值，低于此值的像素被视为内容（非白色区域）
    threshold = 200
    content_mask = image_array < threshold  # 生成内容掩码（True 表示内容区域）

    # 找到内容区域的边界
    coords = np.argwhere(content_mask)  # 获取非白色像素的坐标
    if coords.size == 0:
        # 如果没有找到内容区域，则返回原图像
        print("No content detected, returning original image.")
        return image

    # 获取内容区域的边界
    y_min, x_min = coords.min(axis=0)
    y_max, x_max = coords.max(axis=0)

    # 裁剪图像
    cropped_image = image.crop((x_min, y_min, x_max, y_max))

    # 增加一定的边距，避免公式贴边过紧
    border = 10  # 边距大小
    cropped_image = ImageOps.expand(cropped_image, border=border, fill="white")

    return cropped_image
   
def delete_irrelevant(full_text):
    """
    删除文档中影响选择题识别的干扰内容
    """
    # 删除内容：“原子量：H1   O16   Mg24    Al27    Cl 35.5    Ca40  Fe56    Zn 65”，从而防止干扰
    full_text = re.sub(r"原子量.*?\n", "\n", full_text, flags=re.DOTALL)
    # 如果 full_text 中出现 "第Ⅱ卷"，删除其后的内容，只保留选择题
    match = re.search(r"(?<![\u4e00-\u9fff,，。])第Ⅱ卷", full_text)
    if match:
        full_text = full_text[:match.start()]  # 截取 "第Ⅱ卷" 之前的内容
        
    # 如果 full_text 中出现 "第二部分"，删除其后的内容，只保留选择题
    match = re.search(r"(?<![\u4e00-\u9fff,，。])第二部分", full_text)
    if match:
        full_text = full_text[:match.start()]  # 截取 "第二部分" 之前的内容
  
    # 如果 full_text 中出现 "二、非选择题"，删除其后的内容，只保留选择题
    match = re.search(r"(?<![\u4e00-\u9fff])二、非选择题", full_text)
    if match:
        full_text = full_text[:match.start()]  # 截取 "二、非选择题" 之前的内容
        
    # 如果 full_text 中出现 "三、非选择题"，删除其后的内容，只保留选择题
    match = re.search(r"(?<![\u4e00-\u9fff])三、非选择题", full_text)
    if match:
        full_text = full_text[:match.start()]  # 截取 "三、非选择题" 之前的内容
        
    # 检查 full_text 中是否存在 "一、选择题"
    match = re.search(r"一、选择题", full_text)
    if match:
        # 截取 "一、选择题" 及其后面的内容
        full_text = full_text[match.start():]
        
    # 检查 full_text 中是否存在 "一、单项选择题"
    match = re.search(r"一、单项选择题", full_text)
    if match:
        # 截取 "一、单项选择题" 及其后面的内容
        full_text = full_text[match.start():]
          
    # 检查 full_text 中是否存在 "第I卷"
    match = re.search(r"(?<![\u4e00-\u9fff])第I卷", full_text)
    if match:
        # 截取 "第I卷" 及其后面的内容
        full_text = full_text[match.start():]
        
    # 检查 full_text 中是否存在 "第一部分"
    match = re.search(r"(?<![\u4e00-\u9fff,，。])第一部分", full_text)
    if match:
        # 截取 "第一部分" 及其后面的内容
        full_text = full_text[match.start():]
    
    return full_text

def find_answer(doc, questions, text):
    """
    提取选择题的答案
    """
    # answer_found代表是否按照某种模式已经提取到答案
    answer_found = 0
    
    # 匹配表格中的答案
    table_answers = {}
    for table in doc.tables:
        rows = len(table.rows)
        # 保证表格至少有两行，且行数是偶数（题号和答案成对出现）
        if rows >= 2 and rows % 2 == 0:
            for i in range(0, rows, 2):  # 每次处理两行
                question_numbers = [cell.text.strip() for cell in table.rows[i].cells]  # 奇数行作为题号
                answers = [cell.text.strip() for cell in table.rows[i + 1].cells]       # 偶数行作为答案

                # 将题号与答案对应起来
                for question, answer in zip(question_numbers, answers):
                    if question.isdigit() and re.fullmatch(r'[A-D]+', answer):  # 确保题号是数字，答案为ABCD
                        table_answers[question] = answer
                    
    for question_data in questions:
        index = question_data.get("index")
        if index in table_answers:
            question_data["answer"] = table_answers[index]
            answer_found = 1
            
    # 匹配表格中的答案 1. C
    if answer_found == 0:
        table_answers = {}
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # 匹配单元格中的题号和答案，例如 "1．C"
                    matches = re.findall(r'(\d+)[．.、\s]+([A-D]+)', cell_text)
                    for number, answer in matches:
                        table_answers[number] = answer  # 保存题号和答案

        # 将表格中的答案更新到 questions
        for question_data in questions:
            index = question_data.get("index")
            if index in table_answers:
                question_data["answer"] = table_answers[index]
                answer_found = 1
                    
    
    question_pattern = re.compile(
        r"(?<=[\n \t])"                        # 断言题目前面必须是换行符、空格或制表符
        r"(?P<question>\d+[．.、].{0,300}?)"           # 匹配题目开头，如 "17. 一定质量的..."
        r"(?:([AＡ][．.、\s\u200b]|[\(（]A[\)）])\s*(?P<A>.{0,100}?))"             # 匹配 A 选项
        r"(?:([BＢ][．.、\s\u200b]|[\(（]B[\)）])\s*(?P<B>.{0,100}?))"             # 匹配 B 选项
        r"(?:([CＣ][．.、\s\u200b]|[\(（]C[\)）])\s*(?P<C>.{0,100}?))"             # 匹配 C 选项
        r"(?:([DＤ][．.、\s\u200b]|[\(（]D[\)）])\s*?(?P<D>.{0,100}?))"             # 匹配 D 选项
        r"(?=\n|\f|\d+[．.、][\u4e00-\u9fa5])",     # 断言 D 选项后面是换行符、换页符、题号或中文汉字，但不包含这些内容
        re.DOTALL                           # 允许匹配跨行内容
    )
    
    # 匹配连续的答案部分，形如：
    # 1.C 2.D 3.A 或 1. C 2. D 3. A 或 1、C 2、D 3、A 等形式
    # 14.A,C,D.  15.A,B,D.  16.A,C
    # 6．B、C 　　　7．B 　　　8．A、D
    if answer_found == 0: 
        matches = re.findall(r'(((?<![A-D][.、]\s)(?<![A-D][.．、])\b\d+\b[．.、\s\u3000]+[A-D、.．,，]+\s*){2,})', text)
        # print(matches)
        for match in matches:
            # 提取每个答案（数字和字母）
            answers = re.findall(r'(?<![\u4e00-\u9fa5=.、])\b(\d+)\b[．.、\s\u3000]+([A-D、.．,，]+)', match[0])
            for number, answer in answers:
                # 遍历题目，找到对应的题号并更新答案
                formatted_answer = answer.replace("、", "").replace(".", "").replace("．", "").replace(",", "").replace("，", "").strip()
                for question_data in questions:
                    if question_data.get("index") == number:
                        question_data["answer"] = formatted_answer
                        print("number is: "+number+" answer is: "+formatted_answer)
                        break
                    
            print("答案形式：1.A 2.B")
            answer_found = 1
    
    if answer_found == 0: 
        matches = re.findall(r'((?:\d+[．.、\s\u3000]+\([A-D、,，]+\)\s*){2,})', text)
        for match in matches:
            # 提取每个答案（数字和字母）
            answers = re.findall(r'(\d+)[．.、\s\u3000]+\(([A-D、,，]+)\)', match)
            for number, answer in answers:
                # 遍历题目，找到对应的题号并更新答案
                formatted_answer = answer.replace("、", "").replace(",", "").replace("，", "").strip()
                for question_data in questions:
                    if question_data.get("index") == number:
                        question_data["answer"] = formatted_answer
                        break
            print("答案形式：1.(A) 2.(B)")
            answer_found = 1
            
    # 匹配形如 "1-10 ABCDBCAADB" 的紧凑答案格式
    if answer_found == 0:
        compact_matches = re.findall(r'(\d+)[\-—](\d+)\s+([A-D]+)', text)
        for compact_match in compact_matches:
            start, end, answers = compact_match
            for i, answer in enumerate(answers):
                for question_data in questions:
                    if question_data.get("index") == str(int(start)+i):
                        question_data["answer"] = answer
                        answer_found = 1
                        print("答案形式：1-10 ABCDBCAADB")
                        break
                    
    # 匹配题目后紧跟着 “故选：”
    if answer_found == 0:
        for idx, paragraph in enumerate(doc.paragraphs):
            if re.search(r'故选[:：]\s*([A-D]+)', paragraph.text.strip()):
                # 提取答案内容
                answer_match = re.search(r'故选[:：]\s*([A-D]+)', paragraph.text.strip())
                if answer_match:
                    answer = answer_match.group(1)  # 提取答案
                    answer_found = 1  # 标记找到答案
                    # 拼接从当前段落向前的所有文本
                    combined_text = "\n".join([p.text.strip() for p in doc.paragraphs[:idx+1]])
                    # 查找所有匹配项，从后向前寻找最近的匹配
                    matches = list(question_pattern.finditer(combined_text))
                    if matches:
                        # 获取最后一个匹配项（离“故选：”最近的匹配）
                        question_match = matches[-1]
                        number = (re.match(r'^(\d+)[．.、]', question_match.group(1))).group(1)
                        for question_data in questions:
                            if question_data.get("index") == number and not question_data.get("answer"):
                                question_data["answer"] = answer
                                break
                            
    # 匹配题目后紧跟着【答案】
    if answer_found == 0:
        for idx, paragraph in enumerate(doc.paragraphs):
            if re.match(r'【答案】\s*([A-D]+)', paragraph.text.strip()):
                # 提取答案内容
                answer_match = re.match(r'【答案】\s*([A-D]+)', paragraph.text.strip())
                if answer_match:
                    answer = answer_match.group(1)  # 提取答案
                    answer_found = 1  # 标记找到答案   
                    # 拼接从当前段落向前的所有文本
                    combined_text = "\n".join([p.text.strip() for p in doc.paragraphs[:idx+1]])
                    # 查找所有匹配项，从后向前寻找最近的匹配
                    matches = list(question_pattern.finditer(combined_text))
                    if matches:
                        # 获取最后一个匹配项（离【答案】最近的匹配）
                        question_match = matches[-1]
                        print("question_match is: "+question_match.group(1))
                        number = (re.match(r'^(\d+)[．.、]', question_match.group(1))).group(1)
                        print("number is: "+ number)
                        for question_data in questions:
                            if question_data.get("index") == number and not question_data.get("answer"):
                                question_data["answer"] = answer
                                break
                
    # 匹配类似 “16. 答案： A”
    if answer_found == 0:
        for paragraph in doc.paragraphs:
            if re.match(r'(\d+)[.．]\s*答案：?s*([A-D]+)', paragraph.text.strip()):
                # 提取答案内容
                answer_match = re.match(r'(\d+)[.．]\s*答案：?s*([A-D]+)', paragraph.text.strip())
                if answer_match:
                    index = answer_match.group(1)
                    answer = answer_match.group(2)
                    for question_data in questions:
                        if question_data.get("index") == index:
                            question_data["answer"] = answer
                            break
                    print("答案形式： 答案：")
                    answer_found = 1
                
                
    # 匹配形如 “1、D
    #          （答案内容。。。）
    #          2、A
    #          （答案内容。。。）”
    if answer_found == 0:
        matches = re.findall(r'\n(\d+)[.、\s]+([A-D]+)\s*\n', text)
        for number, answer in matches:
            # 遍历题目，找到对应的题号并更新答案
            for question_data in questions:
                if question_data.get("index") == number:
                    question_data["answer"] = answer
                    print("答案形式：一行一个答案")
                    break
                answer_found = 1
    
def clean_question(question_data):
    """
    删除question中的冗余内容
    """
    # 检查每个 question 条目，删除从某数字到下一个数字的内容
    question_text = question_data.get("question", "")
    # 匹配类似 "数字." 的模式
    current_match = re.match(r'^(\d+)[．.、]', question_text)
    next_match = re.search(r'\n(\d+)[．.、]', question_text)
    
    if current_match and next_match:
        current_num = int(current_match.group(1))
        next_num = int(next_match.group(1))
        print("current is:"+str(current_num)+"next is:"+str(next_num))
        # 如果后一个数字是前一个数字 + 1
        if next_num == current_num + 1:
            # 删除从当前数字到下一个数字之间的内容
            pattern = rf"{current_num}[．.].*?{next_num}[．.]"
            question_text = re.sub(pattern, f"{next_num}.", question_text, flags=re.DOTALL)

        # 更新清理后的 question
        question_data["question"] = question_text
        
def clean_options(questions):
    # 要检查的键
    keys_to_check = ["A", "B", "C", "D"]

    # 遍历每个条目
    for question_data in questions:
        for key in keys_to_check:
            if key in question_data:
                value = question_data[key].strip()
                # 删除以 "）" 开头或以 "（" 结尾的括号
                if value.startswith("）"):
                    value = value[1:]  # 去掉开头的 "）"
                if value.endswith("（"):
                    value = value[:-1]  # 去掉结尾的 "（"
                # 更新条目的值
                question_data[key] = value
    
def add_exam_name(questions):
    """
    增加试卷名
    """
    exam_name = args.docx_name # 提取文件名（去掉路径）
    exam_name = os.path.basename(exam_name)
    exam_name = re.sub(r'真题.*', '', exam_name)
    exam_name = re.sub(r'试题.*', '', exam_name)
    exam_name = re.sub(r'（解析版）.*', '', exam_name)
    exam_name = re.sub(r'（含解析版）.*', '', exam_name)
    for question_data in questions:
        question_data['exam'] = exam_name
        
def save_to_json(output_json_path, questions):
    """
    将生成的条目保存到.json文件，有两种模式：追加和新建
    """
    if args.new == "off":
        # 如果文件已存在，读取其内容
        if os.path.exists(output_json_path):
            with open(output_json_path, "r", encoding="utf-8") as f:
                try:
                    existing_data = json.load(f)  # 读取现有内容
                except json.JSONDecodeError:  # 如果文件为空或内容无效
                    existing_data = []
        else:
            existing_data = []

        # 确保现有内容是列表类型，合并新的内容
        if isinstance(existing_data, list):
            existing_data.extend(questions)  # 将新内容追加到列表
        else:
            raise ValueError("Existing JSON content is not a list. Cannot append.")

        # 保存合并后的内容
        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(existing_data, f, ensure_ascii=False, indent=4)
    else:    
        # 保存为 JSON 文件
        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(questions, f, ensure_ascii=False, indent=4)

def parse_relationships(docx_zip):
    """
    解析 document.xml.rels 文件，建立 rId 与 Target 文件路径的映射关系。
    """
    relationships = {}
    with docx_zip.open('word/_rels/document.xml.rels') as rels_file:
        # 解析 XML
        tree = ET.parse(rels_file)
        root = tree.getroot()
        for rel in root:
            rId = rel.attrib.get("Id")
            target = rel.attrib.get("Target")
            if rId and target:
                relationships[rId] = target
    return relationships

def extract_high_res_image_from_docx(docx_path, rId, relationships):
    """
    从解压后的 Word 文件中提取高分辨率 WMF 图像，并转换为 PNG。
    """
    # 获取 rId 对应的 Target 路径
    target_path = relationships.get(rId)
    print("target path is: "+target_path)
    if not target_path:
        print(f"No target found for rId: {rId}")
        return None

    # 确保路径是相对路径，拼接完整路径
    if not target_path.startswith('/'):
        target_path = f"word/{target_path}"

    # 从 ZIP 文件中提取目标文件
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        if target_path in docx_zip.namelist():
            with docx_zip.open(target_path) as image_file:
                image_data = image_file.read()
                if target_path.endswith(".wmf"):
                    wmf_path = f"./wmf_images/{rId}.wmf"
                    with open(wmf_path, "wb") as wmf_file:
                        wmf_file.write(image_data)
                    output_path = f"./png_images/{rId}.png"
                    output_dir = f"./png_images/"
                    # 构造命令
                    command = [
                        "libreoffice", "--headless", "--convert-to", "png", wmf_path, "--outdir", output_dir
                    ]
                    
                    # 执行命令
                    result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                    return PILImage.open(output_path)
                
                    # image = PILImage.open(output_path)
                    # # 裁剪图像（去除空白区域）
                    # cropped_image = crop_image(image)
                    # # 保存裁剪后的图像
                    # cropped_image.save(f"./png_new_images/{rId}.png")
                    # return cropped_image   
                elif target_path.endswith('.png'):
                    print("png formula")
                    png_path = f"./png_images/{rId}.png"
                    # 保存 PNG 文件
                    with open(png_path, "wb") as png_file:
                        png_file.write(image_data)
                    return PILImage.open(png_path) 
        else:
            print(f"Target path {target_path} not found in ZIP.")
    return None

def extract_formula_from_picture(run, dotx_path, relationships):
    """
    提取段落中的公式图片并识别为 LaTeX。
    """
    run_xml = run.element  # 获取当前运行对象的 XML 元素
    print("picture formula")
    # print(run_xml.xml)
    rId = None
    root = etree.fromstring(run_xml.xml)
    if "<w:object" in run_xml.xml and args.latex == "on":
        print("w:object")
        # 解析 XML 内容
        # 找到 <w:object> 元素
        w_object = root.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
        if w_object is not None:
            v_shape = w_object.find('.//{urn:schemas-microsoft-com:vml}shape') # 找到 <v:shape> 元素
            if v_shape is not None:
                v_imagedata = v_shape.find('.//{urn:schemas-microsoft-com:vml}imagedata') # 找到 <v:imagedata> 元素
                if v_imagedata is not None:
                    rId = v_imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id') # 获取 r:id 属性值
    
    if not rId:
        namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "v": "urn:schemas-microsoft-com:vml",
        }

        # 找到 <w:pict>
        w_pict = root.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
        if w_pict is not None:
            print("<w:pict> found")
            # 找到 <v:shape>
            v_shape = w_pict.find('.//{urn:schemas-microsoft-com:vml}shape')
            if v_shape is not None:
                print("<v:shape> found")
                # 找到 <v:imagedata>
                v_imagedata = v_shape.find('.//{urn:schemas-microsoft-com:vml}imagedata')
                if v_imagedata is not None:
                    print("<v:imagedata> found")
                    rId = v_imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
    if rId:
        # 使用 rId 提取高分辨率图像
        img = extract_high_res_image_from_docx(docx_path, rId, relationships)
        file_path = f"./png_images/{rId}.png"            

        if img:
            # 使用 SimpleTex 的 API 识别公式
            SIMPLETEX_UAT="x97YHMaxT4hl1kbcvKkAHbQqZGR0HDL0rBAZfmqScLusUcO74sXCCOIsNfqO3PgM"
            # api_url="https://server.simpletex.cn/api/latex_ocr"  # 标准模型接口地址
            api_url="https://server.simpletex.cn/api/latex_ocr_turbo"  # 轻量级模型接口地址
            data = { } # 请求参数数据（非文件型参数），视情况填入，可以参考各个接口的参数说明
            header={ "token": SIMPLETEX_UAT } # 鉴权信息，此处使用UAT方式
            file=[("file",(file_path,open(file_path, 'rb')))] # 请求文件,字段名一般为file
            res = requests.post(api_url, files=file, data=data, headers=header) # 使用requests库上传文件
            print(res.text)
            if "err_info" in json.loads(res.text)['res']:
                return ""
            else:
                content = json.loads(res.text)['res']['latex']
                return content
        else:
            print(f"Could not find image for rId: {rId}")
    return ""

def extract_questions_and_answer_from_docx(docx_path, output_json_path):
    # 打开 Word 文档
    doc = Document(docx_path)
    
    # 解压文档，解析 rId 和 Target 的映射关系
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        relationships = parse_relationships(docx_zip)

    # 匹配的正则表达式
    question_pattern = re.compile(
        r"(?<=[\n \t])"                        # 断言题目前面必须是换行符、空格或制表符
        r"(?P<question>\d+[．.、].{0,300}?)"           # 匹配题目开头，如 "17. 一定质量的..."
        r"(?:([AＡ][．.、\s\u200b]|[\(（]A[\)）])\s*(?P<A>.{0,100}?))"             # 匹配 A 选项
        r"(?:([BＢ][．.、\s\u200b]|[\(（]B[\)）])\s*(?P<B>.{0,100}?))"             # 匹配 B 选项
        r"(?:([CＣ][．.、\s\u200b]|[\(（]C[\)）])\s*(?P<C>.{0,100}?))"             # 匹配 C 选项
        r"(?:([DＤ][．.、\s\u200b]|[\(（]D[\)）])\s*?(?P<D>.{0,100}?))"             # 匹配 D 选项
        r"(?=\n|\f|\d+[．.、][\u4e00-\u9fa5])",     # 断言 D 选项后面是换行符、换页符、题号或中文汉字，但不包含这些内容
        re.DOTALL                           # 允许匹配跨行内容
    )
    
    # 遍历段落，提取full_text
    full_text = ""
    # 转换文档中的上标和下标，并提取文本
    for paragraph in doc.paragraphs:
        paragraph_text = ""
        # 遍历段落中的每个 run
        minus = 0
        for run in paragraph.runs:
            paragraph_text += run.text  # 普通文本直接添加
        full_text += paragraph_text + "\n"  # 添加段落并换行

    # 删除文档中影响选择题识别的干扰内容
    original_text = full_text
    full_text = delete_irrelevant(full_text)
    
    # 提取选择题内容
    questions = []
    for match in question_pattern.finditer(full_text):
        question_data = match.groupdict()
        print("question_data[question] is: "+ repr(question_data["question"]))

        # 处理题目段落，提取公式并拼接
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.startswith(question_data["question"][:10].strip()):
                question_start = i
                break
        question_end = 0
        for i in range(question_start, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip().endswith(question_data["question"][-10:].strip()):
                question_end = i 
                break    
        print("\n")        
                
        question_text = ""
        # print("question_start is: " + str(question_start))
        # print("question_end is: " + str(question_end))
        if question_start > question_end:
            question_end = question_start
        
        start = 1
        for i in range(question_start, question_end+1):
            for run in doc.paragraphs[i].runs:
                if run.text.strip():  # 普通文本
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            question_text += f"^{run.text}" 
                        else:
                            question_text += run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            question_text += f"_{run.text}" 
                        else:
                            question_text += run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    else:
                        question_text += run.text  # 普通文本直接添加
                    start = 0
                else:  # 检测并处理公式图片
                    if start == 0:
                        formula_content = extract_formula_from_picture(run, docx_path, relationships)
                        if formula_content:
                            question_text += f" {formula_content} "  # 用 LaTeX 公式替代图片
            question_text += "\n"
        
        question_data["question"] = question_text
        
        clean_question(question_data)
        
        print("question part is: "+question_data["question"])
        print(repr(match.group(0)))
        question_data['index'] = (re.match(r'^(\d+)[．.、]', question_data["question"][:10])).group(1)
      
        # 针对前半部分是试卷，后半部分是试卷+答案的情况
        if any(q['index'] == question_data['index'] for q in questions):
            print(f"Question with index {question_data['index']} already exists. Skipping...")
            continue

        #找到option_paragraph的真正起点
        start_1 = 1
        for j in range(1,len(doc.paragraphs)-question_start):
            if doc.paragraphs[question_start+j].text.replace(" ", "").replace("\t", "").startswith(("A","Ａ","(A)","（A）")):
                start_1 = question_start+j
                break
        option_paragraph1 = doc.paragraphs[start_1]
        option_paragraph2 = None
        for start_2 in range(start_1 + 1, len(doc.paragraphs)):
            if doc.paragraphs[start_2].text.strip():  # 检查段落是否不为空
                option_paragraph2 = doc.paragraphs[start_2]
                break
        for start_3 in range(start_2 + 1, len(doc.paragraphs)):
            if doc.paragraphs[start_3].text.strip():  # 检查段落是否不为空
                option_paragraph3 = doc.paragraphs[start_3]
                break
        for start_4 in range(start_3 + 1, len(doc.paragraphs)):
            if doc.paragraphs[start_4].text.strip():  # 检查段落是否不为空
                option_paragraph4 = doc.paragraphs[start_4]
                break
        results_A = ""
        results_B = ""
        results_C = ""
        results_D = ""
        print("op1: "+option_paragraph1.text)
        print("op2: "+option_paragraph2.text)
        print("op3: "+option_paragraph3.text)
        print("op4: "+option_paragraph4.text)

        #四个选项各占一行
        if option_paragraph1.text.lstrip(" \t").startswith(("A","Ａ","(A)","（A）")) and option_paragraph2.text.lstrip(" \t").startswith(("B","Ｂ","(B)","（B）")): 
            print("[[[Situation 1]]]")
            option_count=0
            minus = 0
            for i,run in enumerate(option_paragraph1.runs):
                if run.text.strip(): # 如果是纯文本，考虑是否存在上下标需要处理
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            result = f"^{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            result = f"_{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    else:
                        result = run.text.lstrip("．.、\t")  # 普通文本直接添加
                else: #不是文本就需要处理图像
                    result = extract_formula_from_picture(run, docx_path, relationships).lstrip("．.、")
                
                if re.search(r"\(A\)|（A）", result) and option_count == 0: 
                    option_count+=1
                    if len(result) > 3: 
                        results_A += result[3:]                        
                elif re.search(r"[AＡ]", result) and option_count == 0: 
                    option_count+=1
                    if len(result) > 2: #考虑到有时 “C．甲、丁的种群数量下降，丙的种群数量增加”会被解析为一个完整的run
                        results_A += result[2:]

                elif option_count==1:
                    results_A += result
            minus = 0
            for i,run in enumerate(option_paragraph2.runs):
                if run.text.strip(): # 如果是纯文本，考虑是否存在上下标需要处理
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            result = f"^{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            result = f"_{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    else:
                        result = run.text.lstrip("．.、\t")  # 普通文本直接添加
                else: #不是文本就需要处理图像
                    result = extract_formula_from_picture(run, docx_path, relationships).lstrip("．.、")

                if re.search(r"\(B\)|（B）", result) and option_count == 1:
                    option_count+=1
                    if len(result) > 3:
                        results_B += result[3:]
                elif re.search(r"[BＢ]", result) and option_count == 1:
                    option_count+=1
                    if len(result) > 2:
                        results_B += result[2:]
                elif option_count==2:
                    results_B += result
            minus = 0
            for i,run in enumerate(option_paragraph3.runs):
                if run.text.strip(): # 如果是纯文本，考虑是否存在上下标需要处理
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            result = f"^{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            result = f"_{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    else:
                        result = run.text.lstrip("．.、\t")  # 普通文本直接添加
                else: #不是文本就需要处理图像
                    result = extract_formula_from_picture(run, docx_path, relationships).lstrip("．.、")

                if re.search(r"\(C\)|（C）", result) and option_count == 2:
                    option_count+=1
                    if len(result) > 3:
                        results_C += result[3:]
                elif re.search(r"[CＣ]", result) and option_count == 2:
                    option_count+=1
                    if len(result) > 2:
                        results_C += result[2:]
                elif option_count==3:
                    results_C += result
            minus = 0
            for i,run in enumerate(option_paragraph4.runs):
                if run.text.strip(): # 如果是纯文本，考虑是否存在上下标需要处理
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            result = f"^{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            result = f"_{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    else:
                        result = run.text.lstrip("．.、\t")  # 普通文本直接添加
                else: #不是文本就需要处理图像
                    result = extract_formula_from_picture(run, docx_path, relationships).lstrip("．.、")

                if re.search(r"\(D\)|（D）", result) and option_count == 3:
                    option_count+=1
                    if len(result) > 3:
                        results_D += result[3:]
                elif re.search(r"[DＤ]", result) and option_count == 3:
                    option_count+=1
                    if len(result) > 2:
                        results_D += result[2:]
                elif option_count==4:
                    results_D += result
        #四个选项占两行
        elif option_paragraph1.text.lstrip(" \t").startswith(("A","Ａ","(A)","（A）")) and option_paragraph2.text.lstrip(" \t").startswith(("C","Ｃ","(C)","（C）")):
            print("[[[Situation 2]]]") 
            option_count=0
            minus = 0
            for i,run in enumerate(option_paragraph1.runs):
                if run.text.strip(): # 如果是纯文本，考虑是否存在上下标需要处理
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            result = f"^{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            result = f"_{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    else:
                        result = run.text.lstrip("．.、\t")  # 普通文本直接添加
                else: #不是文本就需要处理图像
                    result = extract_formula_from_picture(run, docx_path, relationships).lstrip("．.、")

                if (result.strip().startswith(("A","Ａ")) and option_count == 0) or \
                    (result.strip().startswith(("B","Ｂ")) and option_count == 1):
                    if re.search(r"[AＡ]", result) and re.search(r"[BＢ]", result): #考虑：“A. 一直变小 B. 一直变大”
                        # 定义正则表达式匹配整个模式
                        pattern = r"[AＡ][．.、\s](.*?)[BＢ][．.、\s](.*)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)

                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_A = match_option.group(1).strip()  # A. 和 B. 之间的内容
                            results_B = match_option.group(2).strip()  # B. 后面的内容
                            option_count += 2
                    else:
                        option_count += 1
                        if len(result.strip()) > 2:
                            if option_count==1:
                                results_A += result.strip()[2:]
                            else:
                                results_B += result.strip()[2:]
                elif result.strip().startswith(("(A)","（A）")) or result.strip().startswith(("(B)","（B）")):
                    print(result)
                    if re.search(r"\(A\)|（A）", result) and re.search(r"\(B\)|（B）", result): #考虑：“A. 一直变小 B. 一直变大”
                        # 定义正则表达式匹配整个模式
                        print("A and B")
                        pattern = r"(?:\(A\)|（A）)(.*?)(?:\(B\)|（B）)(.*)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)

                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_A = match_option.group(1).strip()  # A. 和 B. 之间的内容
                            results_B = match_option.group(2).strip()  # B. 后面的内容
                    else:
                        option_count+=1
                        if len(result.strip()) > 3:
                            if option_count==1:
                                results_A += result.strip()[3:]
                            else:
                                results_B += result.strip()[3:]
                else:
                    if result.endswith(("B","Ｂ")) and len(result.strip()) > 1 and option_count==1:  #可能遇到 <w:t>＝25cm/s，向左传播         B．</w:t>
                        results_A += result[:-1]
                        option_count += 1
                    elif result.endswith(("B.", "Ｂ.", "B．", "Ｂ．")) and len(result.strip()) > 2 and option_count==1:                       
                        results_A += result[:-2]                              
                        option_count += 1
                    elif result.endswith("(B)") and len(result.strip()) > 3 and option_count==1:
                        results_A += result[:-3]                            
                        option_count += 1                            
                    else:
                        print(result)
                        if option_count==1:
                            results_A += result
                        elif option_count==2:
                            results_B += result
            minus = 0
            for i,run in enumerate(option_paragraph2.runs):
                if run.text.strip(): # 如果是纯文本，考虑是否存在上下标需要处理
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            result = f"^{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            result = f"_{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣":
                            minus = 1
                    else:
                        result = run.text.lstrip("．.、\t")  # 普通文本直接添加
                else: #不是文本就需要处理图像
                    result = extract_formula_from_picture(run, docx_path, relationships).lstrip("．.、")

                if (result.strip().startswith(("C","Ｃ")) and option_count == 2) or \
                    (result.strip().startswith(("D","Ｄ")) and option_count == 3):
                    if re.search(r"[CＣ]", result) and re.search(r"[DＤ]", result): #考虑：“C. 先变小后变大 D. 先变大后变小”
                        # 定义正则表达式匹配整个模式
                        pattern = r"[CＣ][．.、\s](.*?)[DＤ][．.、\s](.*)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)

                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_C = match_option.group(1).strip()  # C. 和 D. 之间的内容
                            results_D = match_option.group(2).strip()  # D. 后面的内容
                            option_count += 2
                    else:
                        option_count+=1
                        if len(result.strip()) > 2:
                            if option_count==3:
                                results_C += result.strip()[2:]
                            else:
                                results_D += result.strip()[2:]
                elif result.strip().startswith(("(C)","（C）")) or result.strip().startswith(("(D)", "（D）")):
                    if re.search(r"\(C\)|（C）", result) and re.search(r"\(D\)|（D）", result): #考虑：“C. 先变小后变大 D. 先变大后变小”
                        # 定义正则表达式匹配整个模式
                        pattern = r"(?:\(C\)|（C）)(.*?)(?:\(D\)|（D）)(.*)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)

                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_C = match_option.group(1).strip()  # C. 和 D. 之间的内容
                            results_D = match_option.group(2).strip()  # D. 后面的内容
                    else:
                        option_count+=1
                        if len(result.strip()) > 3:
                            if option_count==3:
                                results_C += result.strip()[3:]
                            else:
                                results_D += result.strip()[3:]
                else:
                    if (result.endswith(("C", "Ｃ")) and len(result.strip()) > 1 and option_count==2) or \
                    (result.endswith(("D", "Ｄ")) and len(result.strip()) > 1 and option_count==3): 
                        if option_count==2:
                            results_B += result[:-1]  
                        elif option_count==3:
                            results_C += result[:-1]
                        option_count += 1
                    elif (result.endswith(("C.","Ｃ.", "C．","Ｃ．")) and len(result.strip()) > 2 and option_count==2) or \
                    (result.endswith(("D.","Ｄ.", "D．","Ｄ．")) and len(result.strip()) > 2 and option_count==3):
                        if option_count==2:
                            results_B += result[:-2]                               
                        elif option_count==3:
                            results_C += result[:-2]
                        option_count += 1
                    elif (result.endswith("(C)") and len(result.strip()) > 3 and option_count==2) or \
                    (result.endswith("(D)") and len(result.strip()) > 3 and option_count==3):
                        if option_count==2:
                            results_B += result[:-3]
                        elif option_count==3:
                            results_C += result[:-3]                               
                        option_count += 1                               
                    else:
                        if option_count==3:
                            results_C += result
                        elif option_count==4:
                            results_D += result                        
        else: #四个选项在同一行
            print("[[[Situation 3]]]")
            option_count=0
            minus = 0
            for i,run in enumerate(option_paragraph1.runs):
                if run.text.strip(): # 如果是纯文本，考虑是否存在上下标需要处理
                    if run.font.superscript:  # 如果是上标
                        if minus == 0:
                            result = f"^{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    elif run.font.subscript:  # 如果是下标
                        if minus == 0:
                            result = f"_{run.text}" 
                        else:
                            result = run.text
                            minus = 0
                        if run.text == "﹣" or run.text == "－":
                            minus = 1
                    else:
                        result = run.text.lstrip("．.、\t") # 普通文本直接添加
                else: #不是文本就需要处理图像
                    result = extract_formula_from_picture(run, docx_path, relationships).lstrip("．.、")
                    print(result)
                if re.search(r"\(A\)|（A）", result) or re.search(r"\(B\)|（B）", result) or re.search(r"\(C\)|（C）", result) or re.search(r"\(D\)|（D）", result):
                    #全是文本，一整行被解析成一个run
                    if re.search(r"\(A\)|（A）", result) and re.search(r"\(B\)|（B）", result) and re.search(r"\(C\)|（C）", result) and re.search(r"\(D\)|（D）", result):
                        # 定义正则表达式匹配整个模式
                        pattern = r"(?:\(A\)|（A）)(.*?)(?:\(B\)|（B）)(.*?)(?:\(C\)|（C）)(.*?)(?:\(D\)|（D）)(.*)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)

                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_A = match_option.group(1).strip()  # A. 和 B. 之间的内容
                            results_B = match_option.group(2).strip()  # B. 和 C. 之间的内容
                            results_C = match_option.group(3).strip()  # C. 和 D. 之间的内容
                            results_D = match_option.group(4).strip()  # D. 后面的内容
                    elif re.search(r"\(A\)|（A）", result) and re.search(r"\(B\)|（B）", result):
                        # 定义正则表达式匹配整个模式
                        pattern = r"(?:\(A\)|（A）)(.*?)(?:\(B\)|（B）)(.*)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)
                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_A = match_option.group(1).strip()  # A. 和 B. 之间的内容
                            results_B = match_option.group(2).strip()  # B. 和 C. 之间的内容
                        option_count = 2
                    elif re.search(r"\(B\)|（B）", result) and re.search(r"\(C\)|（C）", result):
                        # 定义正则表达式匹配整个模式
                        pattern = r"(?:\(B\)|（B）)(.*?)(?:\(C\)|（C）)(.*?)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)
                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_B = match_option.group(1).strip()  # B. 和 C. 之间的内容
                            results_C = match_option.group(2).strip()  # C. 和 D. 之间的内容
                        option_count = 3
                    elif re.search(r"\(C\)|（C）", result) and re.search(r"\(D\)|（D）", result):
                        # 定义正则表达式匹配整个模式
                        pattern = r"(?:\(C\)|（C）)(.*?)(?:\(D\)|（D）)(.*?)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)
                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_C = match_option.group(1).strip()  # C. 和 D. 之间的内容
                            results_D = match_option.group(2).strip()  
                        option_count = 4
                    else:
                        option_count += 1

                        if len(result.strip()) > 3 and result.strip().startswith(("(A)","（A）", "(B)","（B）", "(C)","（C）", "(D)","（D）")):
                            if option_count==1:
                                results_A += result.strip()[3:]
                            elif option_count==2:
                                results_B += result.strip()[3:]
                            elif option_count==3:
                                results_C += result.strip()[3:]
                            else:
                                results_D += result.strip()[3:]   
                elif (re.search(r"(^|[^a-zA-Z0-9_{}.]\^)[AＡ]([^a-zA-Z0-9_{}\^]|$)", result) and option_count == 0) or \
                     (re.search(r"(^|[^a-zA-Z0-9_{}.\^])[BＢ]([^a-zA-Z0-9_{}\^]|$)", result) and option_count == 1) or \
                     (re.search(r"(^|[^a-zA-Z0-9_{}.\^])[CＣ]([^a-zA-Z0-9_{}\^]|$)", result) and option_count == 2) or \
                     (re.search(r"(^|[^a-zA-Z0-9_{}.\^])[DＤ]([^a-zA-Z0-9_{}\^]|$)", result) and option_count == 3):
                    #全是文本，一整行被解析成一个run
                    print("A or B or C or D")
                    print(result)
                    if re.search(r"[AＡ]", result) and re.search(r"[BＢ]", result) and re.search(r"[CＣ]", result) and re.search(r"[DＤ]", result):
                        # 定义正则表达式匹配整个模式
                        pattern = r"[AＡ][．.、\s](.*?)[BＢ][．.、\s](.*?)[CＣ][．.、\s](.*?)[DＤ][．.、\s](.*)"

                        # 使用 re.search 寻找第一次匹配
                        match_option = re.search(pattern, result)

                        # 提取各部分内容并去掉多余空白
                        if match_option:
                            results_A = match_option.group(1).strip()  # A. 和 B. 之间的内容
                            results_B = match_option.group(2).strip()  # B. 和 C. 之间的内容
                            results_C = match_option.group(3).strip()  # C. 和 D. 之间的内容
                            results_D = match_option.group(4).strip()  # D. 后面的内容
                    else:
                        if len(result.strip()) > 2 and result.strip().startswith(("A","Ａ", "B","Ｂ", "C","Ｃ", "D","Ｄ")):
                            option_count += 1
                            if result.endswith(("B","Ｂ", "C","Ｃ", "D","Ｄ")): #考虑到可能出现 <w:t xml:space="preserve">    B．0        C．</w:t>
                                if option_count==1:
                                    results_A += result.strip()[2:-1]
                                elif option_count==2:
                                    results_B += result.strip()[2:-1]
                                elif option_count==3:
                                    results_C += result.strip()[2:-1]
                                else:
                                    results_D += result.strip()[2:-1]  
                                option_count += 1 
                            elif result.endswith(("B.","Ｂ.", "C.","Ｃ.", "D.","Ｄ.", "B．","Ｂ．", "C．","Ｃ．", "D．","Ｄ．")):  
                                if option_count==1:
                                    results_A += result.strip()[2:-2]
                                elif option_count==2:
                                    results_B += result.strip()[2:-2]
                                elif option_count==3:
                                    results_C += result.strip()[2:-2]
                                else:
                                    results_D += result.strip()[2:-2]  
                                option_count += 1                                    
                            else:
                                if option_count==1:
                                    print("A")
                                    results_A += result.strip()[2:]
                                elif option_count==2:
                                    print("B")
                                    results_B += result.strip()[2:]
                                elif option_count==3:
                                    print("C")
                                    results_C += result.strip()[2:]
                                else:
                                    print("D")
                                    results_D += result.strip()[2:]
                        elif len(result.strip()) > 2:
                            print("result len > 2 and result is: " + result)
                            if option_count==1 and re.search(r"(^|[^a-zA-Z0-9_{}.\^])[BＢ]([^a-zA-Z0-9_{}\^]|$)", result):
                                match_before_B = re.search(r"^(.*?)\s*[BＢ]", result.strip())
                                if match_before_B:
                                    print("before B: "+match_before_B.group(1).strip())
                                    results_A += match_before_B.group(1).strip() 
                                option_count += 1
                                match_after_B = re.search(r"[BＢ][．.、\s](.*)", result.strip())
                                if match_after_B:
                                    print("after B: "+match_after_B.group(1).strip())
                                    results_B += match_after_B.group(1).strip()
                            
                            if option_count==2 and re.search(r"(^|[^a-zA-Z0-9_{}.\^])[CＣ]([^a-zA-Z0-9_{}\^]|$)", result):
                                match_before_C = re.search(r"^(.*?)\s*[CＣ]", result.strip())
                                if match_before_C:
                                    print("before C: "+match_before_C.group(1).strip())
                                    results_B += match_before_C.group(1).strip() 
                                option_count += 1
                                match_after_C = re.search(r"[CＣ][．.、\s](.*)", result.strip())
                                if match_after_C:
                                    print("after C: "+match_after_C.group(1).strip())
                                    results_C += match_after_C.group(1).strip()
                            
                            if option_count==3 and re.search(r"(^|[^a-zA-Z0-9_{}.\^])[DＤ]([^a-zA-Z0-9_{}\^]|$)", result):
                                match_before_D = re.search(r"^(.*?)\s*[DＤ]", result.strip())
                                if match_before_D:
                                    print("before D: "+match_before_D.group(1).strip())
                                    results_C += match_before_D.group(1).strip() 
                                option_count += 1
                                match_after_D = re.search(r"[DＤ][．.、\s](.*)", result.strip())
                                if match_after_D:
                                    print("after D: "+match_after_D.group(1).strip())
                                    results_D += match_after_D.group(1).strip()                    
                        else:
                            option_count += 1    
                elif option_count==1:
                    results_A += result
                elif option_count==2:
                    results_B += result
                elif option_count==3:
                    results_C += result
                else:
                    results_D += result
        question_data["A"] = results_A
        question_data["B"] = results_B
        question_data["C"] = results_C
        question_data["D"] = results_D
        print(">>>>>>>>>>")  
            
        # 添加识别到的问题
        questions.append(question_data)
        
    clean_options(questions)   
     
    find_answer(doc, questions, original_text)
 
    add_exam_name(questions)
            
    # 追加到 JSON 文件
    save_to_json(output_json_path, questions)


# 命令行入口
if __name__ == "__main__":
    # 定义命令行参数
    parser = argparse.ArgumentParser(description="从 Word 文档中提取选择题和答案并保存为 JSON 格式")
    parser.add_argument("--docx_name", type=str, help="Word 文档名")
    parser.add_argument("--json_name", type=str, help="json 文件名")
    parser.add_argument("--new", type=str, help="控制是否追加到之前的输出上，on代表新建.json文件，off代表追加到.json文件后")
    parser.add_argument("--latex", type=str, default="on", help="是否开启 Simpletex 图片识别")
    args = parser.parse_args()
    
    docx_path = f"./{args.docx_name}.docx"  
    output_json_path = f"./{args.json_name}.json"  
    extract_questions_and_answer_from_docx(docx_path, output_json_path)
    print(f"完成.json生成，文件已保存到 {args.json_name}.json")
    